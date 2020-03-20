from docx import Document
from docx.shared import Inches
import csv

data =[]

with open('mahdata.csv') as fileko:
    reader = csv.reader(fileko)
    for row in reader:
        data.append(row)






document = Document()

document.add_heading('Ampogi ni al', 1)


table = document.add_table(rows=1, cols=6)
table.style = 'TableGrid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Province_State'
hdr_cells[1].text = 'Country_Region'
hdr_cells[2].text = 'LastUpdate'
hdr_cells[3].text = 'Confirmed'
hdr_cells[4].text = 'Deaths'
hdr_cells[5].text = 'Recovered'
rownum=0
for c1,c2,c3,c4,c5,c6 in data:
    row_cells = table.add_row().cells
    if rownum != 0:
        row_cells[0].text = c1
        row_cells[1].text = c2
        row_cells[2].text = c3
        row_cells[3].text = c4
        row_cells[4].text = c5
        row_cells[5].text = c6
        print(f'Adding row {rownum}')
    rownum += 1

document.add_page_break()

document.save('output.docx')